import sys
import os

try:
    import docx
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 1. 에듀집 탑재용 체크리스트 생성
def create_eduzip_checklist():
    doc = Document()
    doc.add_heading('학습지원 소프트웨어 필수기준 체크리스트(공급자용)', 0)

    # 개요
    doc.add_heading('□ 제품/서비스 개요', level=1)
    table1 = doc.add_table(rows=3, cols=4)
    table1.style = 'Table Grid'
    
    table1.cell(0, 0).text = '제품/서비스명'
    table1.cell(0, 1).text = '우리 반 칭찬 스티커판'
    table1.cell(0, 2).text = '공급자'
    table1.cell(0, 3).text = '중앙대학교사범대학부속초등학교 교사 박세원'
    
    table1.cell(1, 0).text = '접속경로'
    table1.cell(1, 1).text = 'https://github.com/sewon-ship-it/20260627compliment (현재 로컬/Github 버전)'
    table1.cell(1, 1).merge(table1.cell(1, 3))
    
    table1.cell(2, 0).text = '주요 내용 및 기능·특장점'
    table1.cell(2, 1).text = '◦ 교사용 칭찬 스티커 누적 관리 웹앱\n◦ 교실 브라우저 로컬 저장소를 활용하여 외부 서버 데이터 전송 원천 차단\n◦ 별도 회원가입 없이 즉시 사용 가능한 교육용 편의 도구'
    table1.cell(2, 1).merge(table1.cell(2, 3))

    doc.add_paragraph()

    # 개인정보보호 기준 충족여부
    doc.add_heading('□ 개인정보보호 기준 충족여부', level=1)
    table2 = doc.add_table(rows=10, cols=6)
    table2.style = 'Table Grid'
    
    # Headers
    headers = ['선정기준', '세부 내용', '충족', '미충족', '해당없음', '증빙']
    for i, h in enumerate(headers):
        table2.cell(0, i).text = h
        
    data = [
        ('1. 최소처리 원칙 준수', '1-1. 개인정보가 최소한으로 수집되는가?', '■', '□', '□', '(개인정보 처리방침 제3조 등) 당해 학년도 종료 시 파기 원칙'),
        ('1. 최소처리 원칙 준수', '1-2. 개인정보 수집·이용 목적이 기재되어 있는가?', '■', '□', '□', '(개인정보 처리방침 제1조 등) 교사의 학급 운영 목적 기재'),
        ('1. 최소처리 원칙 준수', '1-3. 개인정보 수집항목, 보유기간 등이 기재되어 있는가?', '■', '□', '□', '(개인정보 처리방침 제2조, 제3조 등)'),
        ('2. 개인정보 안전조치 의무', '2-1. 개인정보 안전성 확보에 필요한 조치 사항이 기재되어 있는가?', '■', '□', '□', '(개인정보 처리방침 제6조 등) 로컬 브라우저 보안 샌드박스 내부 저장'),
        ('3. 열람/정정/삭제/처리정지 절차', '3-1. 이용자에게 언제든지 자신의 정보를 열람·정정·삭제·처리정지를 요구할 수 있는 절차가 안내되어 있는가?', '■', '□', '□', '(개인정보 처리방침 제7조 등)'),
        ('4. 만14세 미만 아동의 개인정보 보호', '4-1. 만 14세 미만 아동의 경우 법정대리인 동의 등 아동의 개인정보 보호를 위한 절차가 마련되어 있는가?', '■', '□', '□', '(이용약관 제6조 등) 교사를 통한 보호자 사전 동의 절차 명시'),
        ('5. 보호책임자/제3자제공/위탁 등', '5-1. 개인정보 보호책임자 관련 정보가 안내되어 있는가?', '■', '□', '□', '(개인정보 처리방침 제8조 등) 책임자: 교사 박세원'),
        ('5. 보호책임자/제3자제공/위탁 등', '5-2. 개인정보 제3자 제공에 관한 정보가 기재되어 있는가? (필요시)', '■', '□', '□', '(개인정보 처리방침 제5조 등) 외부 제공 없음 명시'),
        ('5. 보호책임자/제3자제공/위탁 등', '5-3. 개인정보 위·수탁관계에 관한 정보가 기재되어 있는가? (필요시)', '■', '□', '□', '(해당없음) 외부 위탁 없음'),
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, text in enumerate(row_data):
            table2.cell(row_idx, col_idx).text = text
            
    # Merge column 1 cells where criteria is the same
    table2.cell(1, 0).merge(table2.cell(3, 0))
    table2.cell(7, 0).merge(table2.cell(9, 0))
            
    doc.add_paragraph('\n□ 작성일 및 문의처\n작성일: 2026. 06. 27.\n작성자: 박세원 교사')
    
    doc.save('에듀집_탑재용_체크리스트(우리반칭찬스티커판).docx')

# 2. 학교용 필수기준 체크리스트 생성
def create_school_checklist():
    doc = Document()
    doc.add_heading('학습지원 소프트웨어 선정기준 체크리스트 (학교용)', 0)

    # 필수기준
    doc.add_heading('1. 필수기준', level=1)
    table1 = doc.add_table(rows=10, cols=6)
    table1.style = 'Table Grid'
    
    headers = ['선정기준', '세부 내용', '충족', '미충족', '해당없음', '증빙자료']
    for i, h in enumerate(headers):
        table1.cell(0, i).text = h
        
    data1 = [
        ('1. 최소처리 원칙 준수', '1-1. 개인정보가 최소한으로 수집되는가?', '■', '□', '□', '개인정보 수집을 최소화하고 있습니다. (로컬 데이터)'),
        ('1. 최소처리 원칙 준수', '1-2. 개인정보 수집·이용 목적이 기재되어 있는가?', '■', '□', '□', '수집 및 이용 목적을 약관 제1조에 명확히 기재하였습니다.'),
        ('1. 최소처리 원칙 준수', '1-3. 개인정보 수집항목, 보유기간 등이 기재되어 있는가?', '■', '□', '□', '수집 항목과 보유 기간을 상세히 안내하고 있습니다.'),
        ('2. 개인정보 안전조치 의무', '2-1. 개인정보 안전성 확보에 필요한 조치 사항이 기재되어 있는가?', '■', '□', '□', '기술적·관리적 보호 조치(로컬 저장, 통신 차단)를 포함합니다.'),
        ('3. 열람/정정/삭제/처리정지 절차', '3-1. 이용자에게 언제든지 자신의 정보를 열람·정정·삭제·처리정지를 요구할 수 있는 절차가 안내되어 있는가?', '■', '□', '□', '정보주체의 권리 행사 절차를 보장하고 있습니다.'),
        ('4. 만14세 미만 아동의 개인정보 보호', '4-1. 만 14세 미만 아동의 경우 법정대리인 동의 등 아동의 개인정보 보호를 위한 절차가 마련되어 있는가?', '■', '□', '□', '법정대리인 동의 절차를 준수합니다.'),
        ('5. 보호책임자/제3자제공/위탁 등', '5-1. 개인정보 보호책임자 관련 정보가 안내되어 있는가?', '■', '□', '□', '개인정보 보호책임자를 명시하고 있습니다. (박세원 교사)'),
        ('5. 보호책임자/제3자제공/위탁 등', '5-2. 개인정보 제3자 제공에 관한 정보가 기재되어 있는가? (필요시)', '□', '□', '■', '개인정보를 제 3 자 에 게 제공하지 않습니다.'),
        ('5. 보호책임자/제3자제공/위탁 등', '5-3. 개인정보 위·수탁관계에 관한 정보가 기재되어 있는가? (필요시)', '□', '□', '■', '클라우드 인프라 활용 및 외부 위탁 내용이 없습니다 (로컬 전용).'),
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        for col_idx, text in enumerate(row_data):
            table1.cell(row_idx, col_idx).text = text
            
    table1.cell(1, 0).merge(table1.cell(3, 0))
    table1.cell(7, 0).merge(table1.cell(9, 0))

    doc.add_paragraph()

    # 선택기준
    doc.add_heading('2. 선택기준', level=1)
    table2 = doc.add_table(rows=6, cols=6)
    table2.style = 'Table Grid'
    
    headers2 = ['선정기준', '세부 내용', '충족', '미충족', '해당없음', '증빙자료']
    for i, h in enumerate(headers2):
        table2.cell(0, i).text = h
        
    data2 = [
        ('1. 교육목표 및 학생특성 적합성', '1-1. 수업 목표와 학생의 학습 수준에 적합한 내용과 기능을 제공하는가?', '■', '□', '□', '학생의 긍정 행동 강화 및 인성 교육 목표에 부합합니다.'),
        ('2. 콘텐츠 품질 및 안전성', '2-1. 학습 콘텐츠가 정확하고 신뢰할 수 있으며, 학생 연령에 적합·안전한가?', '■', '□', '□', '서버 통신 없는 로컬 환경 설계로 데이터 유출 가능성을 원천 차단했습니다.'),
        ('3. 사용 환경 적합성', '3-1. 학교의 기기·네트워크 환경에서 모든 학생이 안정적으로 사용할 수 있는가?', '■', '□', '□', '웹 표준 기술(HTML/React) 기반으로 모든 기기 및 브라우저 환경에 최적화되었습니다.'),
        ('4. 접근성 및 사용성', '4-1. 교사와 학생이 필요한 기능과 자료에 쉽게 접근하고 활용할 수 있는가?', '■', '□', '□', '사용자 중심의 직관적인 디자인을 갖췄습니다.'),
        ('5. 서비스 운영 및 지원체계', '5-1. 이용 안내, 기술 지원, 문의 대응 등 서비스 지원 체계를 갖추고 있는가?', '■', '□', '□', '명확한 이용약관 및 보호책임자 문의 채널을 운영 중입니다.'),
    ]

    for row_idx, row_data in enumerate(data2, start=1):
        for col_idx, text in enumerate(row_data):
            table2.cell(row_idx, col_idx).text = text

    doc.save('학교용_필수기준_체크리스트(우리반칭찬스티커판).docx')

if __name__ == '__main__':
    create_eduzip_checklist()
    create_school_checklist()
    print("DOCX files generated successfully.")
